import re
import collections
import decimal

from docx import Document
from pdfplumber import utils
import pdfplumber

from splitter.read_pdf.table import TableFinder
from splitter.read_pdf.my_utils import extract_text
from splitter.base.models import FileDatabase, TextData, TableData

pdf_type = pdfplumber.pdf.PDF

re_page_number = re.compile(r'[—-]+\s?\d+\s?[—-]+')  # pdf页码


class RePdfBlack(object):

    def __init__(self):
        self.re_dpf_blank1 = re.compile(r'(\d)\s([^0-9])')
        self.re_dpf_blank2 = re.compile(r'([^0-9])\s(\d)')

    def sub(self, repl, string):
        data = self.re_dpf_blank1.sub(repl, string)
        return self.re_dpf_blank2.sub(repl, data)


re_pdf_blank = RePdfBlack()


class PDFDealBase(object):
    """PDF中的各类格式数据处理，泛化数据处理存储等模块"""

    @property
    def settings(self):
        return {}

    def read(self):
        """获取每一段的内容"""
        yield

    def to_word(self, document):
        return document


class PDFDealWords(PDFDealBase):
    """PDF中的文本处理"""

    def __init__(self, page):
        self.page = page

    def deal_str(self, data):
        """去除PDF读取出的数据中的页码（仅限固定格式，目前纯数字的不行）
        和数字与字符之间的空格（发现）数字与中文衔接处会多一个空格，原因是PDF是按位算字符的"""
        data = re_page_number.sub(r'', data)  # 去除页码
        data = re_pdf_blank.sub(r'\1\2', data)
        return data

    @property
    def settings(self):
        """配置文本处理配置，主要是判断同行同列时的偏差阈值，默认为3"""
        hs = []
        for char in self.page.chars:
            h = char['y1'] - char['y0']
            hs.append(h)
        try:  # 求文字高度的众数，取众数值的一半作为同行误差的阈值
            y_tolerance = collections.Counter(hs).most_common(1)[0][0]
            y_tolerance = float(y_tolerance)//2
        except:
            y_tolerance = utils.DEFAULT_Y_TOLERANCE
        # by_center = utils.cluster_objects(words, lambda x: (x["x0"] + x["x1"])/2, 1)
        return {
            "x_tolerance": utils.DEFAULT_X_TOLERANCE,
            "y_tolerance": y_tolerance,  # utils.DEFAULT_Y_TOLERANCE,
        }

    def read(self):
        # yield self.page.extract_text(**self.settings)
        text = extract_text(self.page.chars, **self.settings)
        if text is None:
            yield ""  # 之前返回None现在返回空字符
        else:
            for paragraph in text.split('\n'):
                yield paragraph


# class PDFDealTable(PDFDealBase):
#     """PDF中的表格处理"""
#
#     def __init__(self, table):
#         """新版的代码可以直接处理表，而且未采用修改源码的方式而是重写了表格处理模块"""
#         self.table = table
#
#     @property
#     def settings(self):
#         """
#         配置表格处理配置项，后期根据实际情况配置
#         目前传入的已是处理好的表格，因此该方法不再必要
#         """
#         return {}
#
#     def read(self):
#         yield self.table.extract()
#
#     def to_word(self, document):
#         for lines in self.read():
#             table = document.add_table(rows=0, cols=len(lines[0]))  # 插入表格
#             table.style = 'Table Grid'
#             for line in lines:
#                 row_cells = table.add_row().cells
#                 for n, column in enumerate(line, 0):
#                     if column is None:
#                         column = ''
#                     row_cells[n].text = str(column)
#         return document


class ReadPDF(FileDatabase):
    """读取PDF内容，目前仅读取文本内容和表格"""
    def __init__(self, pdf: pdf_type, table_settings=None):
        super(ReadPDF, self).__init__()
        self.pdf = pdf
        self.table_settings = table_settings or {
            "vertical_strategy": "lines",
            "horizontal_strategy": "lines",
        }
        self.initialize()

    @classmethod
    def open(cls, file_path: str, table_settings=None, password="", **kwargs):
        pdf = pdfplumber.open(file_path, password=password, **kwargs)
        return cls(pdf, table_settings=table_settings)

    def close(self):
        self.pdf.close()

    def initialize(self):
        for page_num, page in enumerate(self.pdf.pages, 1):  # 读取每一页
            # 确定文本宽度（对于有页眉页码的可以通过该步骤去除）
            size = self.header_footer_size(page.bbox[1], page.bbox[3])
            page = page.crop((page.bbox[0], page.bbox[1] + size, page.bbox[2], page.bbox[3] - size))
            # 将页面内容根据类型剪切成一段段的（文本、表格）
            self.crop_bbox(page, page_num)

    def crop_bbox(self, page, page_num, *args, **kwargs):
        """将页面内容根据表格和文本处理成有序队列,返回文本、表格两种格式的对象组成的队列"""
        # """将页面内容根据表格和文本处理成有序队列，标记数据位置与类型"""
        tables = TableFinder(page, settings=self.table_settings).find_table().tables
        words, top = [], page.bbox[1]
        for table in tables:
            bbox = (page.bbox[0], top, page.bbox[2], table.bbox[1])
            for paragraph in PDFDealWords(page.crop(bbox)).read():
                self.add(TextData(paragraph, page=page_num))
            self.add(TableData(table.extract(*args, **kwargs), page=page_num))
            top = table.bbox[3]
        bbox = (page.bbox[0], top, page.bbox[2], page.bbox[3])
        for paragraph in PDFDealWords(page.crop(bbox)).read():
            self.add(TextData(paragraph, page=page_num))
        return self

    def header_footer_size(self, top, bottom, size=25.4, h=297):
        """A4纸默认页眉页脚为25.4mm,文本高度是297"""
        size = decimal.Decimal(size)
        h = decimal.Decimal(h)
        return size/h * (bottom - top)

    def to_word(self, file_path):
        """读取PDF内容，目前仅读取文本内容和表格"""
        document = Document()  # 创建一个word对象
        new_page_num = 1
        for paragraph in self.doc:
            paragraph.to_word(document)
            if paragraph.page != new_page_num:
                document.add_page_break()
                new_page_num = paragraph.page
        document.save(file_path)


class ReadPDF2(FileDatabase):
    """
    读取PDF内容，仅读取文本内容,不处理其中的表格，表格会当成文本处理
    该方法导致表格后接的段落与表格被认定为同一行，因此弃用此方法，此方法的效果在
    ReadPDF中的get_all_str中实现
    """
    def __init__(self, pdf: pdf_type):
        super(ReadPDF2, self).__init__()
        self.pdf = pdf
        self.initialize()

    @classmethod
    def open(cls, file_path: str, password="", **kwargs):
        pdf = pdfplumber.open(file_path, password=password, **kwargs)
        return cls(pdf)

    def close(self):
        self.pdf.close()

    def initialize(self):
        for page_num, page in enumerate(self.pdf.pages, 1):  # 读取每一页
            # 确定文本宽度（对于有页眉页码的可以通过该步骤去除）
            size = self.header_footer_size(page.bbox[1], page.bbox[3])
            page = page.crop((page.bbox[0], page.bbox[1] + size, page.bbox[2], page.bbox[3] - size))
            for paragraph in PDFDealWords(page).read():
                self.add(TextData(paragraph, page=page_num))

    def header_footer_size(self, top, bottom, size=25.4, h=297):
        """A4纸默认页眉页脚为25.4mm,文本高度是297"""
        size = decimal.Decimal(size)
        h = decimal.Decimal(h)
        return size/h * (bottom - top)

    def to_word(self, file_path):
        """读取PDF内容，目前仅读取文本内容和表格"""
        document = Document()  # 创建一个word对象
        new_page_num = 1
        for paragraph in self.doc:
            paragraph.to_word(document)
            if paragraph.page != new_page_num:
                document.add_page_break()
                new_page_num = paragraph.page
        document.save(file_path)
