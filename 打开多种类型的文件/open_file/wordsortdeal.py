from collections import Iterable

from docx import Document

from open_file.settings import logger


class WordSortDealBase(object):

    def __init__(self, fname):
        self.fname = fname
        self.docx = Document(self.fname)

    def extract_text(self, paragraphs_numbers=None):
        '''
        :param paragraphs_numbers: default None, all text
        '''
        if paragraphs_numbers is None:
            text = '\n'.join(paragraph.text for paragraph in self.docx.paragraphs)
        elif isinstance(paragraphs_numbers, int):
            text = self.docx.paragraphs[paragraphs_numbers].text
        elif isinstance(paragraphs_numbers, Iterable):
            try:
                text = '\n'.join(self.docx.paragraphs[i].text for i in paragraphs_numbers)
            except Exception:
                logger.debug('page_numbers must be an array of Numbers or Numbers or None.')
                raise Exception('page_numbers must be an array of Numbers or Numbers or None.')
        else:
            logger.debug('page_numbers must be an array of Numbers or Numbers or None.')
            raise Exception('page_numbers must be an array of Numbers or Numbers or None.')
        # \u3000是空格，用于首行缩进，这里以此为特征，认为其代表新段落的开始
        # 用于弥补当前处理的Word文档中常常存在无换行的问题，该问题导致匹配标题序号时无法做到精确处理
        # 使用此方法可以尽可能的防止此情况发生,之后\n会写入正则中，因此不建议修改回来
        return text


def deal_spetial2enter(text):
    text = text.replace('\u3000', '\n')
    text = text.replace('\xa0', '\n')
    return text

